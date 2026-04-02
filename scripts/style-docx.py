#!/usr/bin/env python3
"""
Post-process Word documents to add professional styling.
Adds table borders, code block backgrounds, and other formatting.

Usage:
    python3 scripts/style-docx.py input.docx output.docx
"""

import sys
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_table_borders(table):
    """Add professional borders and shading to tables."""
    # Set table borders
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Add table borders
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # Border width
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')  # Black borders
        tblBorders.append(border)
    
    tblPr.append(tblBorders)
    
    # Add alternating row shading (header + alternating rows)
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            # Header row (first row) - dark gray background
            if i == 0:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'D9D9D9')  # Light gray
                cell._element.get_or_add_tcPr().append(shading)
                # Make header text bold
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            # Alternating rows - very light gray
            elif i % 2 == 0:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'F2F2F2')  # Very light gray
                cell._element.get_or_add_tcPr().append(shading)

def style_code_blocks(doc):
    """Add gray background to code blocks and prevent page breaks within them."""
    for paragraph in doc.paragraphs:
        # Check if paragraph is a code block (Pandoc uses "Source Code" style)
        if paragraph.style.name == 'Source Code':
            # Add gray background
            pPr = paragraph._element.get_or_add_pPr()
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'E8E8E8')  # Light gray background
            pPr.append(shading)
            
            # Prevent page breaks within code blocks
            # Add "keep lines together" property
            keepLines = OxmlElement('w:keepLines')
            pPr.append(keepLines)
            
            # Add "keep with next" to prevent splitting from following paragraph
            keepNext = OxmlElement('w:keepNext')
            pPr.append(keepNext)
            
            # Add padding/spacing
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.right_indent = Inches(0.25)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            
            # Ensure monospace font
            for run in paragraph.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)

def keep_headings_with_content(doc):
    """Prevent page breaks between headings and their following content."""
    for paragraph in doc.paragraphs:
        # Check if paragraph is a heading
        if paragraph.style.name.startswith('Heading'):
            pPr = paragraph._element.get_or_add_pPr()
            
            # Add "keep with next" to prevent heading from being separated from content
            keepNext = OxmlElement('w:keepNext')
            pPr.append(keepNext)
            
            # Prevent page break before heading (optional, for cleaner layout)
            # pageBreakBefore = OxmlElement('w:pageBreakBefore')
            # pageBreakBefore.set(qn('w:val'), '0')
            # pPr.append(pageBreakBefore)

def style_document(input_path, output_path):
    """Apply professional styling to Word document."""
    print(f"📄 Loading document: {input_path}")
    doc = Document(input_path)
    
    # Style all tables
    print(f"📊 Styling {len(doc.tables)} tables...")
    for table in doc.tables:
        add_table_borders(table)
    
    # Style code blocks
    print("💻 Styling code blocks...")
    style_code_blocks(doc)
    
    # Keep headings with their content
    print("📑 Preventing heading orphans...")
    keep_headings_with_content(doc)
    
    # Save styled document
    print(f"💾 Saving styled document: {output_path}")
    doc.save(output_path)
    print("✅ Done!")

def main():
    if len(sys.argv) != 3:
        print("Usage: python3 scripts/style-docx.py input.docx output.docx")
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2]
    
    try:
        style_document(input_path, output_path)
    except Exception as e:
        print(f"❌ Error: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()